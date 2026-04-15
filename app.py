import streamlit as st
from docxtpl import DocxTemplate
from docx import Document
import io
import zipfile
import pandas as pd
from datetime import date
import re
import subprocess
import tempfile
import os

# ─── Page Config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="DocsFill",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── Custom CSS ───────────────────────────────────────────────────────────────
st.markdown("""
<style>
/* ── Base & Typography ────────────────────────────────────────────────────── */
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600&family=Sarabun:wght@300;400;500;600&display=swap');

html, body, [class*="css"] {
    font-family: 'Sarabun', 'Inter', sans-serif;
}

/* ── Background ───────────────────────────────────────────────────────────── */
.stApp {
    background-color: #FAF9F6;
}

/* ── Sidebar ──────────────────────────────────────────────────────────────── */
[data-testid="stSidebar"] {
    background-color: #FFFFFF;
    border-right: 1px solid #EEEBE4;
}
[data-testid="stSidebar"] .stMarkdown h1,
[data-testid="stSidebar"] .stMarkdown h2,
[data-testid="stSidebar"] .stMarkdown h3,
[data-testid="stSidebar"] .stMarkdown h4 {
    color: #2C2C2C;
    font-weight: 600;
}
[data-testid="stSidebar"] .stRadio label {
    font-size: 0.9rem;
    color: #555;
}

/* ── Main Header ──────────────────────────────────────────────────────────── */
.main-header {
    padding: 2rem 0 0.5rem 0;
    border-bottom: 1px solid #EEEBE4;
    margin-bottom: 2rem;
}
.main-header h1 {
    font-size: 1.8rem;
    font-weight: 600;
    color: #1A1A1A;
    margin: 0;
    letter-spacing: -0.02em;
}
.main-header p {
    color: #888;
    font-size: 0.9rem;
    margin: 0.25rem 0 0 0;
}

/* ── Cards ────────────────────────────────────────────────────────────────── */
.card {
    background: #FFFFFF;
    border: 1px solid #EEEBE4;
    border-radius: 12px;
    padding: 1.5rem;
    margin-bottom: 1rem;
}

/* ── Step Indicator ───────────────────────────────────────────────────────── */
.step-container {
    display: flex;
    align-items: center;
    gap: 0;
    margin-bottom: 2rem;
}
.step-item {
    display: flex;
    align-items: center;
    gap: 0.5rem;
    flex: 1;
}
.step-dot {
    width: 28px;
    height: 28px;
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 0.75rem;
    font-weight: 600;
    flex-shrink: 0;
}
.step-dot.done   { background: #2C2C2C; color: #FFFFFF; }
.step-dot.active { background: #C8A96E; color: #FFFFFF; }
.step-dot.idle   { background: #EEEBE4; color: #AAA; }
.step-label {
    font-size: 0.82rem;
    font-weight: 500;
}
.step-label.done   { color: #2C2C2C; }
.step-label.active { color: #C8A96E; }
.step-label.idle   { color: #BBB; }
.step-line {
    height: 1px;
    flex: 0.5;
    background: #EEEBE4;
    margin: 0 0.25rem;
}

/* ── Upload Area ──────────────────────────────────────────────────────────── */
[data-testid="stFileUploader"] {
    background: #FFFFFF;
    border: 1.5px dashed #D9D4CA;
    border-radius: 12px;
    padding: 0.5rem;
    transition: border-color 0.2s;
}
[data-testid="stFileUploader"]:hover {
    border-color: #C8A96E;
}

/* ── Form & Inputs ────────────────────────────────────────────────────────── */
.stTextInput input,
.stTextArea textarea,
.stNumberInput input,
.stDateInput input {
    border: 1px solid #DDD9D0 !important;
    border-radius: 8px !important;
    background: #FFFFFF !important;
    font-family: 'Sarabun', 'Inter', sans-serif !important;
    color: #2C2C2C !important;
    transition: border-color 0.2s !important;
}
.stTextInput input:focus,
.stTextArea textarea:focus,
.stNumberInput input:focus {
    border-color: #C8A96E !important;
    box-shadow: 0 0 0 3px rgba(200,169,110,0.12) !important;
}
.stTextInput label,
.stTextArea label,
.stNumberInput label,
.stDateInput label {
    font-size: 0.82rem !important;
    font-weight: 500 !important;
    color: #555 !important;
    text-transform: uppercase;
    letter-spacing: 0.04em;
}

/* ── Buttons ──────────────────────────────────────────────────────────────── */
.stButton > button[kind="primary"],
.stFormSubmitButton > button[kind="primary"] {
    background: #2C2C2C !important;
    color: #FFFFFF !important;
    border: none !important;
    border-radius: 8px !important;
    font-weight: 500 !important;
    font-size: 0.9rem !important;
    padding: 0.6rem 1.5rem !important;
    transition: background 0.2s, transform 0.1s !important;
    letter-spacing: 0.01em;
}
.stButton > button[kind="primary"]:hover,
.stFormSubmitButton > button[kind="primary"]:hover {
    background: #444 !important;
    transform: translateY(-1px);
}
.stButton > button[kind="secondary"] {
    background: transparent !important;
    border: 1px solid #DDD9D0 !important;
    color: #555 !important;
    border-radius: 8px !important;
}
.stDownloadButton > button {
    background: #F5F0E8 !important;
    color: #2C2C2C !important;
    border: 1px solid #DDD9D0 !important;
    border-radius: 8px !important;
    font-weight: 500 !important;
}
.stDownloadButton > button:hover {
    background: #EDE6D6 !important;
    border-color: #C8A96E !important;
}

/* ── Alerts ───────────────────────────────────────────────────────────────── */
[data-testid="stAlert"] {
    border-radius: 8px !important;
    border-left-width: 3px !important;
}

/* ── Info box ─────────────────────────────────────────────────────────────── */
.var-badge {
    display: inline-block;
    background: #F5F0E8;
    color: #8B6914;
    border: 1px solid #E8DCC8;
    border-radius: 5px;
    padding: 1px 8px;
    font-size: 0.78rem;
    font-family: 'Courier New', monospace;
    margin: 2px;
}

/* ── Divider ──────────────────────────────────────────────────────────────── */
hr {
    border: none !important;
    border-top: 1px solid #EEEBE4 !important;
    margin: 1.5rem 0 !important;
}

/* ── Dataframe ────────────────────────────────────────────────────────────── */
[data-testid="stDataFrame"] {
    border: 1px solid #EEEBE4;
    border-radius: 8px;
    overflow: hidden;
}

/* ── Progress bar ─────────────────────────────────────────────────────────── */
.stProgress > div > div {
    background-color: #C8A96E !important;
    border-radius: 4px;
}

/* ── Expander ─────────────────────────────────────────────────────────────── */
[data-testid="stExpander"] {
    border: 1px solid #EEEBE4 !important;
    border-radius: 10px !important;
    background: #FFFFFF !important;
}

/* ── Hide Streamlit branding ──────────────────────────────────────────────── */
#MainMenu, footer { visibility: hidden; }
</style>
""", unsafe_allow_html=True)

# ─── Helper: ตรวจชื่อตัวแปรแล้วเลือก input widget ที่เหมาะสม ─────────────────
def get_input_widget(var: str, label: str, key: str):
    v = var.lower()
    if re.search(r"date|วันที่|วันเกิด|start_date|end_date|เริ่ม|สิ้นสุด", v):
        val = st.date_input(label, key=key)
        return val.strftime("%d/%m/%Y")
    elif re.search(r"note|detail|description|remark|address|ที่อยู่|หมายเหตุ|รายละเอียด", v):
        return st.text_area(label, key=key, placeholder=f"กรอก {label}", height=80)
    elif re.search(r"email|อีเมล", v):
        return st.text_input(label, key=key, placeholder="example@email.com")
    elif re.search(r"phone|tel|mobile|โทร|เบอร์", v):
        return st.text_input(label, key=key, placeholder="0XX-XXX-XXXX")
    elif re.search(r"amount|price|cost|salary|ราคา|เงิน|ค่า|จำนวนเงิน|budget", v):
        val = st.number_input(label, key=key, min_value=0.0, step=1.0, format="%.2f")
        return f"{val:,.2f}"
    else:
        return st.text_input(label, key=key, placeholder=f"กรอก {label}")

# ─── Helper: สร้าง sample template ───────────────────────────────────────────
@st.cache_data
def create_sample_template() -> bytes:
    doc = Document()
    doc.add_heading("หนังสือรับรองการทำงาน", 0)
    doc.add_paragraph("")
    doc.add_paragraph("วันที่  {{ date }}")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run(
        "หนังสือฉบับนี้ขอรับรองว่า  {{ full_name }}  อายุ  {{ age }}  ปี "
        "ดำรงตำแหน่ง  {{ position }}  ได้ปฏิบัติงานกับ  {{ company_name }} "
        "ตั้งแต่วันที่  {{ start_date }}  จนถึงปัจจุบัน "
        "โดยได้รับเงินเดือน  {{ salary }}  บาทต่อเดือน"
    )
    doc.add_paragraph("")
    doc.add_paragraph("ที่อยู่ปัจจุบัน:  {{ address }}")
    doc.add_paragraph("อีเมล:  {{ email }}")
    doc.add_paragraph("เบอร์โทรศัพท์:  {{ phone }}")
    doc.add_paragraph("")
    doc.add_paragraph("จึงออกหนังสือรับรองฉบับนี้ให้ไว้เพื่อเป็นหลักฐาน")
    doc.add_paragraph("")
    doc.add_paragraph("ลงชื่อ  ____________________________")
    doc.add_paragraph("(  {{ approver_name }}  )")
    doc.add_paragraph("{{ approver_position }}")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ─── Helper: เรนเดอร์เทมเพลต ─────────────────────────────────────────────────
def render_doc(template_bytes: bytes, context: dict) -> bytes:
    doc = DocxTemplate(io.BytesIO(template_bytes))
    doc.render(context)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ─── Helper: แปลง docx → PDF ─────────────────────────────────────────────────
def convert_to_pdf(docx_bytes: bytes) -> bytes:
    tmp_dir = tempfile.mkdtemp()
    tmp_docx = os.path.join(tmp_dir, "input.docx")
    tmp_pdf  = os.path.join(tmp_dir, "input.pdf")
    try:
        with open(tmp_docx, "wb") as f:
            f.write(docx_bytes)
        try:
            from docx2pdf import convert
            convert(tmp_docx, tmp_pdf)
        except Exception:
            soffice_candidates = [
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                "soffice",
            ]
            converted = False
            for soffice in soffice_candidates:
                result = subprocess.run(
                    [soffice, "--headless", "--convert-to", "pdf",
                     "--outdir", tmp_dir, tmp_docx],
                    capture_output=True, timeout=60,
                )
                if result.returncode == 0:
                    converted = True
                    break
            if not converted:
                raise RuntimeError(
                    "ไม่พบ Microsoft Word หรือ LibreOffice\n"
                    "กรุณาติดตั้ง LibreOffice: https://www.libreoffice.org/download/libreoffice-fresh/"
                )
        with open(tmp_pdf, "rb") as f:
            return f.read()
    finally:
        for p in [tmp_docx, tmp_pdf]:
            if os.path.exists(p):
                os.unlink(p)
        os.rmdir(tmp_dir)

# ─── Helper: Step Indicator ───────────────────────────────────────────────────
def show_steps(current: int):
    steps = [
        ("1", "อัปโหลด Template"),
        ("2", "กรอกข้อมูล"),
        ("3", "ดาวน์โหลด"),
    ]
    parts = []
    for i, (num, label) in enumerate(steps, start=1):
        state = "done" if i < current else ("active" if i == current else "idle")
        parts.append(f"""
        <div class="step-item">
            <div class="step-dot {state}">{("✓" if state == "done" else num)}</div>
            <span class="step-label {state}">{label}</span>
        </div>
        """)
        if i < len(steps):
            parts.append('<div class="step-line"></div>')

    st.markdown(
        f'<div class="step-container">{"".join(parts)}</div>',
        unsafe_allow_html=True,
    )

# ─── Sidebar ──────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div style="padding: 0.5rem 0 1rem 0;">
        <div style="font-size: 1.5rem; margin-bottom: 0.25rem;">📄</div>
        <div style="font-size: 1.1rem; font-weight: 600; color: #1A1A1A;">DocsFill</div>
        <div style="font-size: 0.78rem; color: #AAA; margin-top: 0.1rem;">Word Template Filler · v2.0</div>
    </div>
    """, unsafe_allow_html=True)

    st.divider()

    mode = st.radio(
        "โหมดการใช้งาน",
        ["Single — กรอกทีละเอกสาร", "Batch — สร้างจาก CSV"],
        help="Single: กรอกข้อมูลเองทีละชุด | Batch: นำเข้า CSV สร้างหลายไฟล์พร้อมกัน",
        label_visibility="visible",
    )

    st.divider()

    st.markdown('<p style="font-size:0.78rem;font-weight:600;color:#AAA;text-transform:uppercase;letter-spacing:0.06em;margin-bottom:0.5rem;">ตัวอย่าง Template</p>', unsafe_allow_html=True)
    st.caption("หนังสือรับรองการทำงาน พร้อมตัวแปรครบ")
    st.download_button(
        label="ดาวน์โหลด Sample .docx",
        data=create_sample_template(),
        file_name="sample_template.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

    st.divider()

    st.markdown('<p style="font-size:0.78rem;font-weight:600;color:#AAA;text-transform:uppercase;letter-spacing:0.06em;margin-bottom:0.5rem;">รูปแบบ Output</p>', unsafe_allow_html=True)
    output_format = st.radio(
        "เลือก format",
        ["DOCX", "PDF", "DOCX และ PDF"],
        help="PDF ต้องการ Microsoft Word หรือ LibreOffice",
        label_visibility="collapsed",
    )

    st.divider()

    with st.expander("วิธีสร้าง Template"):
        st.markdown("""
ใส่ตัวแปรใน Word ด้วยรูปแบบนี้:

| ตัวแปร | Input |
|--------|-------|
| `{{ name }}` | Text |
| `{{ date }}` | Date Picker |
| `{{ email }}` | Email |
| `{{ phone }}` | เบอร์โทร |
| `{{ address }}` | Textarea |
| `{{ salary }}` | ตัวเลข |
        """)

# ─── Main Header ──────────────────────────────────────────────────────────────
if "Single" in mode:
    st.markdown("""
    <div class="main-header">
        <h1>สร้างเอกสาร</h1>
        <p>อัปโหลด Template แล้วกรอกข้อมูลเพื่อสร้างเอกสาร Word</p>
    </div>
    """, unsafe_allow_html=True)
else:
    st.markdown("""
    <div class="main-header">
        <h1>Batch Generation</h1>
        <p>อัปโหลด Template + CSV เพื่อสร้างเอกสารหลายไฟล์พร้อมกัน</p>
    </div>
    """, unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# MODE: SINGLE
# ══════════════════════════════════════════════════════════════════════════════
if "Single" in mode:

    show_steps(1)

    uploaded_file = st.file_uploader(
        "อัปโหลดไฟล์ Word Template (.docx)",
        type=["docx"],
        help="ไฟล์ต้องมีตัวแปร {{ variable }} อยู่ภายใน",
        label_visibility="collapsed",
    )

    if not uploaded_file:
        st.markdown("""
        <div style="text-align:center; padding: 3rem 0; color: #BBB;">
            <div style="font-size: 2.5rem; margin-bottom: 0.75rem;">☁️</div>
            <div style="font-size: 0.95rem; color: #999;">ลากไฟล์มาวางที่นี่ หรือคลิกเพื่อเลือกไฟล์</div>
            <div style="font-size: 0.8rem; color: #CCC; margin-top: 0.25rem;">รองรับเฉพาะ .docx เท่านั้น</div>
        </div>
        """, unsafe_allow_html=True)

    if uploaded_file:
        try:
            template_bytes = uploaded_file.read()
            doc = DocxTemplate(io.BytesIO(template_bytes))
            variables = doc.get_undeclared_template_variables()

            if not variables:
                st.warning("ไม่พบตัวแปรใดๆ ในเทมเพลตนี้ — ลองดาวน์โหลด Sample Template ใน sidebar")
            else:
                show_steps(2)

                vars_list = sorted(variables)

                # แสดง badge ตัวแปรที่พบ
                badges = "".join([f'<span class="var-badge">{v}</span>' for v in vars_list])
                st.markdown(
                    f'<div style="background:#FDFCF8;border:1px solid #EEEBE4;border-radius:8px;padding:0.75rem 1rem;margin-bottom:1.5rem;">'
                    f'<span style="font-size:0.78rem;color:#AAA;font-weight:600;text-transform:uppercase;letter-spacing:0.05em;">พบ {len(vars_list)} ตัวแปร</span>'
                    f'<div style="margin-top:0.4rem">{badges}</div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )

                with st.form("single_form"):
                    st.markdown('<p style="font-size:0.78rem;font-weight:600;color:#AAA;text-transform:uppercase;letter-spacing:0.06em;margin-bottom:1rem;">กรอกข้อมูล</p>', unsafe_allow_html=True)
                    context: dict = {}

                    if len(vars_list) > 3:
                        col1, col2 = st.columns(2, gap="large")
                        for i, var in enumerate(vars_list):
                            label = var.replace("_", " ").title()
                            with col1 if i % 2 == 0 else col2:
                                context[var] = get_input_widget(var, label, key=f"s_{var}")
                    else:
                        for var in vars_list:
                            label = var.replace("_", " ").title()
                            context[var] = get_input_widget(var, label, key=f"s_{var}")

                    st.markdown("<div style='margin-top:1.5rem'></div>", unsafe_allow_html=True)
                    submitted = st.form_submit_button(
                        "สร้างเอกสาร", use_container_width=True, type="primary"
                    )

                if submitted:
                    empty = [k for k, v in context.items() if not str(v).strip()]
                    if empty:
                        st.error("กรุณากรอกให้ครบ ยังขาด: " + ", ".join(empty))
                    else:
                        with st.spinner("กำลังสร้างเอกสาร..."):
                            docx_output = render_doc(template_bytes, context)

                        show_steps(3)
                        st.success("สร้างเอกสารสำเร็จ")

                        base_name = uploaded_file.name.replace(".docx", "_filled")
                        want_docx = "DOCX" in output_format
                        want_pdf  = "PDF"  in output_format

                        dl_cols = st.columns(2) if (want_docx and want_pdf) else [st.container()]

                        if want_docx:
                            with dl_cols[0]:
                                st.download_button(
                                    label="ดาวน์โหลด DOCX",
                                    data=docx_output,
                                    file_name=f"{base_name}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True,
                                    type="primary",
                                )

                        if want_pdf:
                            with (dl_cols[1] if want_docx else dl_cols[0]):
                                with st.spinner("กำลังแปลงเป็น PDF..."):
                                    try:
                                        pdf_output = convert_to_pdf(docx_output)
                                        st.download_button(
                                            label="ดาวน์โหลด PDF",
                                            data=pdf_output,
                                            file_name=f"{base_name}.pdf",
                                            mime="application/pdf",
                                            use_container_width=True,
                                            type="primary",
                                        )
                                    except RuntimeError as pdf_err:
                                        st.error(f"แปลง PDF ไม่ได้: {pdf_err}")

        except Exception as e:
            st.error(f"ไม่สามารถอ่านไฟล์ได้: {e}")


# ══════════════════════════════════════════════════════════════════════════════
# MODE: BATCH
# ══════════════════════════════════════════════════════════════════════════════
else:
    col1, col2 = st.columns(2, gap="large")
    with col1:
        st.markdown('<p style="font-size:0.78rem;font-weight:600;color:#AAA;text-transform:uppercase;letter-spacing:0.06em;margin-bottom:0.5rem;">1 · Word Template</p>', unsafe_allow_html=True)
        uploaded_template = st.file_uploader(
            "Word Template", type=["docx"], label_visibility="collapsed"
        )
    with col2:
        st.markdown('<p style="font-size:0.78rem;font-weight:600;color:#AAA;text-transform:uppercase;letter-spacing:0.06em;margin-bottom:0.5rem;">2 · ไฟล์ข้อมูล CSV</p>', unsafe_allow_html=True)
        uploaded_csv = st.file_uploader(
            "CSV File",
            type=["csv"],
            help="Row แรกต้องเป็น header ชื่อตรงกับตัวแปรใน Template",
            label_visibility="collapsed",
        )

    # ── อัปโหลด Template แล้ว ยังไม่มี CSV ───────────────────────────────────
    if uploaded_template and not uploaded_csv:
        try:
            tmpl_b = uploaded_template.read()
            doc = DocxTemplate(io.BytesIO(tmpl_b))
            variables = doc.get_undeclared_template_variables()
            if variables:
                badges = "".join([f'<span class="var-badge">{v}</span>' for v in sorted(variables)])
                st.markdown(
                    f'<div style="background:#FDFCF8;border:1px solid #EEEBE4;border-radius:8px;padding:0.75rem 1rem;margin:1rem 0;">'
                    f'<span style="font-size:0.78rem;color:#AAA;font-weight:600;text-transform:uppercase;letter-spacing:0.05em;">พบ {len(variables)} ตัวแปร — ดาวน์โหลด CSV Template แล้วกรอกข้อมูล</span>'
                    f'<div style="margin-top:0.4rem">{badges}</div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )
                sample_df = pd.DataFrame(columns=sorted(variables))
                st.download_button(
                    "ดาวน์โหลด CSV Template",
                    data=sample_df.to_csv(index=False),
                    file_name="data_template.csv",
                    mime="text/csv",
                    use_container_width=True,
                )
                st.dataframe(sample_df, use_container_width=True)
            st.session_state["tmpl_bytes"] = tmpl_b
        except Exception as e:
            st.error(f"อ่านไฟล์ Template ไม่ได้: {e}")

    # ── อัปโหลดครบทั้งคู่ ────────────────────────────────────────────────────
    if uploaded_template and uploaded_csv:
        try:
            tmpl_b = uploaded_template.read()
            df = pd.read_csv(uploaded_csv, dtype=str).fillna("")

            st.markdown("<div style='margin-top:1rem'></div>", unsafe_allow_html=True)

            # Summary row
            c1, c2 = st.columns(2)
            with c1:
                st.markdown(
                    f'<div style="background:#FFFFFF;border:1px solid #EEEBE4;border-radius:10px;padding:1rem 1.25rem;">'
                    f'<div style="font-size:0.75rem;color:#AAA;font-weight:600;text-transform:uppercase;letter-spacing:0.05em;">จำนวนแถว</div>'
                    f'<div style="font-size:1.8rem;font-weight:600;color:#1A1A1A;margin-top:0.1rem;">{len(df)}</div>'
                    f'</div>', unsafe_allow_html=True
                )
            with c2:
                st.markdown(
                    f'<div style="background:#FFFFFF;border:1px solid #EEEBE4;border-radius:10px;padding:1rem 1.25rem;">'
                    f'<div style="font-size:0.75rem;color:#AAA;font-weight:600;text-transform:uppercase;letter-spacing:0.05em;">จำนวนคอลัมน์</div>'
                    f'<div style="font-size:1.8rem;font-weight:600;color:#1A1A1A;margin-top:0.1rem;">{len(df.columns)}</div>'
                    f'</div>', unsafe_allow_html=True
                )

            st.markdown("<div style='margin-top:1rem'></div>", unsafe_allow_html=True)

            with st.expander("ดูตัวอย่างข้อมูล (5 แถวแรก)", expanded=True):
                st.dataframe(df.head(), use_container_width=True)

            st.markdown("<div style='margin-top:1rem'></div>", unsafe_allow_html=True)

            if st.button(
                f"สร้างเอกสารทั้งหมด {len(df)} ไฟล์",
                use_container_width=True,
                type="primary",
            ):
                progress_bar = st.progress(0, text="เริ่มต้น...")
                status_text = st.empty()
                zip_buf = io.BytesIO()

                want_docx = "DOCX" in output_format
                want_pdf  = "PDF"  in output_format

                with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                    for i, row in df.iterrows():
                        context = {col: str(val) for col, val in row.items()}
                        doc_bytes = render_doc(tmpl_b, context)

                        first_col = df.columns[0]
                        safe_name = (
                            str(row[first_col])
                            .replace("/", "-")
                            .replace(" ", "_")
                            .replace("\\", "-")[:50]
                        )
                        base = f"{int(i)+1:03d}_{safe_name}"

                        if want_docx:
                            zf.writestr(f"{base}.docx", doc_bytes)
                        if want_pdf:
                            try:
                                pdf_bytes = convert_to_pdf(doc_bytes)
                                zf.writestr(f"{base}.pdf", pdf_bytes)
                            except RuntimeError:
                                zf.writestr(f"{base}.docx", doc_bytes)

                        pct = (int(i) + 1) / len(df)
                        progress_bar.progress(pct, text=f"สร้างแล้ว {int(i)+1}/{len(df)}")
                        status_text.caption(f"→ {base}")

                zip_buf.seek(0)
                status_text.empty()
                st.success(f"สร้างครบทั้ง {len(df)} เอกสาร")

                st.download_button(
                    label=f"ดาวน์โหลดทั้งหมด ({len(df)} ไฟล์) เป็น ZIP",
                    data=zip_buf,
                    file_name="generated_documents.zip",
                    mime="application/zip",
                    use_container_width=True,
                    type="primary",
                )

        except Exception as e:
            st.error(f"เกิดข้อผิดพลาด: {e}")
